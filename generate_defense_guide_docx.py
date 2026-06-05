from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = Path(r"D:\XM\Food-Court\美食街摊位管理系统_答辩学习指南_标准答题卡.docx")


def set_run_font(run, font_name="宋体", size=Pt(11), bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.35):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_paragraph(doc, text, *, style=None, bold=False, size=Pt(11), align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    set_paragraph_spacing(p)
    return p


def add_bullets(doc, items, indent_level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.74 + indent_level * 0.5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item)
        set_run_font(run, size=Pt(11))
        set_paragraph_spacing(p, after=3)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item)
        set_run_font(run, size=Pt(11))
        set_paragraph_spacing(p, after=3)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    if level == 1:
        p.style = "Heading 1"
        size = Pt(16)
    elif level == 2:
        p.style = "Heading 2"
        size = Pt(14)
    else:
        p.style = "Heading 3"
        size = Pt(12)
    run = p.add_run(text)
    set_run_font(run, font_name="黑体", size=size, bold=True)
    set_paragraph_spacing(p, before=6, after=6, line=1.2)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    header_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        cell = header_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, font_name="黑体", size=Pt(10.5), bold=True)
        shade_cell(cell, "D9EAF7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths and i < len(widths):
            cell.width = widths[i]

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(text))
            set_run_font(run, size=Pt(10.5))
            set_paragraph_spacing(p, after=0, line=1.2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths and i < len(widths):
                cell.width = widths[i]
    doc.add_paragraph("")


def add_question_card(doc, idx, question, answer, points=None):
    add_heading(doc, f"{idx}. {question}", 3)
    add_paragraph(doc, answer)
    if points:
        add_bullets(doc, points)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    if "List Bullet" in styles:
        styles["List Bullet"].font.name = "宋体"
        styles["List Bullet"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        styles["List Bullet"].font.size = Pt(11)
    if "List Number" in styles:
        styles["List Number"].font.name = "宋体"
        styles["List Number"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        styles["List Number"].font.size = Pt(11)

    for heading_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[heading_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.bold = True

    if "Quote" not in styles:
        styles.add_style("Quote", WD_STYLE_TYPE.PARAGRAPH)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("美食街摊位管理系统")
    set_run_font(run, font_name="黑体", size=Pt(24), bold=True, color=(31, 78, 121))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    run = p2.add_run("答辩学习指南")
    set_run_font(run, font_name="黑体", size=Pt(20), bold=True)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(40)
    run = p3.add_run("附：老师高频追问 30 题标准答题卡")
    set_run_font(run, font_name="黑体", size=Pt(16), bold=True, color=(89, 89, 89))

    info_lines = [
        "文档用途：答辩前系统复习、项目理解、演示串讲、老师追问准备",
        "项目技术：Java 17 + Jakarta Servlet 6.0 + JSP 3.1 + JDBC + HikariCP + MySQL 8.0 + Bootstrap 5",
        "项目代码主目录：D:\\XM\\Food-Court\\food-court-management",
        "内容依据：项目源码、SQL 脚本、thesis.md、core_code.md",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_run_font(run, size=Pt(11))
        set_paragraph_spacing(p, after=6)

    doc.add_page_break()


def main():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "一、这份文档怎么用", 1)
    add_paragraph(
        doc,
        "这不是论文摘抄版，而是按照“答辩最容易被问到什么、项目最应该怎么理解、现场最容易卡壳在哪里”的思路整理出来的学习指南。你可以把它当成一份复习剧本：先建立项目总图，再抓核心链路，最后记住高频追问的标准回答口径。"
    )
    add_bullets(
        doc,
        [
            "时间非常紧时，优先看“项目速览”“核心业务主线”“当前不足与改进口径”“老师高频追问 30 题”。",
            "准备演示时，重点熟悉食客下单、跨摊位拆单、摊主接单、管理员统计与租赁审核这四条主线。",
            "准备老师追问时，不要只背优点，必须把明文密码、事务缺失、鉴权分散、会话级购物车等限制点也准备好。老师往往更喜欢你能清楚说明边界和改进计划。",
            "这份材料严格依据当前代码实现整理。没有在代码中出现的功能，不会当作既有成果写进标准答案。"
        ]
    )

    add_heading(doc, "二、项目速览", 1)
    add_paragraph(
        doc,
        "一句话概括：这是一个面向食客、摊主、管理员三类角色的 Java Web 美食街摊位管理系统，覆盖“浏览商品、加入购物车、统一结算、按摊位自动拆单、摊主处理订单、管理员治理平台”的完整业务闭环。"
    )
    add_table(
        doc,
        ["维度", "本项目实际情况"],
        [
            ["核心定位", "美食街/食堂档口/夜市集中经营场景下的轻量级 Web 管理系统"],
            ["角色设计", "食客 DINER、摊主 OWNER、管理员 ADMIN"],
            ["后端架构", "Servlet 控制器 + Service 业务层 + DAO 持久层 + MySQL"],
            ["前端形态", "JSP + JSTL + EL + Bootstrap 5"],
            ["核心业务亮点", "跨摊位购物车统一结算并按 stallId 自动拆单"],
            ["治理能力", "用户状态管理、摊位管理、品类容量、订单监控、租赁审核、合同文本生成"],
            ["工程边界", "课程/毕设级实现，安全、事务、一致性和分布式能力仍有明显提升空间"],
        ],
        widths=[Cm(4), Cm(11.5)]
    )
    add_paragraph(
        doc,
        "如果老师让你先整体介绍项目，你可以按“场景痛点 -> 三类角色 -> 分层架构 -> 核心亮点 -> 当前不足”这五步讲，逻辑最稳。"
    )

    add_heading(doc, "三、答辩开场模板", 1)
    add_heading(doc, "1. 30 秒极简版", 2)
    add_paragraph(
        doc,
        "我的项目是一个基于 Java Web 的美食街摊位管理系统，面向食客、摊主和管理员三类角色。食客可以浏览摊位和商品、管理购物车并下单；摊主可以维护商品和处理订单；管理员可以管理用户、摊位、品类容量和租赁审核。项目最大的业务特点是支持跨摊位统一结算，并能按摊位自动拆单，便于不同摊主独立接单和结算。"
    )
    add_heading(doc, "2. 3 分钟标准版", 2)
    add_numbered(
        doc,
        [
            "项目背景：传统美食街或食堂档口通常依赖线下排队、人工记单、分散统计，跨摊位消费体验差，摊主接单效率低，管理员也难以及时掌握整体经营情况。",
            "项目目标：构建一个基于 B/S 架构的轻量级管理系统，把食客点餐、摊主经营、管理员治理放到同一个 Web 平台中完成。",
            "技术方案：后端采用 Java 17、Jakarta Servlet、JSP、JDBC 和 HikariCP，前端使用 Bootstrap 5，数据库采用 MySQL 8.0。",
            "功能主线：食客端支持注册登录、浏览摊位、商品选购、购物车、优惠券、提交订单和查看订单；摊主端支持看经营数据、管理商品和处理订单；管理员端支持用户管理、摊位管理、订单监控、品类容量维护和租赁审核。",
            "核心亮点：食客可以在一次结算中购买多个摊位的商品，系统在结算时会按摊位进行分组，为每个摊位分别生成订单和订单明细，实现统一支付、分摊位接单。",
            "当前不足：项目目前仍是课程级实现，存在密码明文存储、订单创建未做数据库事务包裹、鉴权逻辑分散在各个 Servlet 中、部分扩展表尚未真正接入业务等问题，后续可以向更完整的工程化版本继续演进。"
        ]
    )

    add_heading(doc, "四、先用什么顺序理解这个项目", 1)
    add_paragraph(
        doc,
        "很多同学答辩时容易出现“代码看过，但讲不成体系”的问题。原因通常不是不会，而是理解顺序错了。正确的理解顺序应该是先业务，再结构，再数据，再关键代码，最后是不足和优化。"
    )
    add_numbered(
        doc,
        [
            "先记三类角色：食客是消费侧，摊主是经营侧，管理员是治理侧。",
            "再记一条主链路：浏览商品 -> 加购物车 -> 统一结算 -> 自动拆单 -> 摊主处理 -> 管理员统计。",
            "再看系统结构：controller 收请求，service 组织业务，dao 负责 SQL，entity 表示对象，model 存放购物车这样的会话模型，util 负责数据库连接等公共能力。",
            "再记数据库：users、stalls、products、orders、order_items、leases 是当前业务的六张核心表，order_number_sequences 和 pickup_number_sequences 用来支持每天的编号递增。",
            "最后背不足：明文密码、缺少统一过滤器、订单创建没有事务、摊主更新订单时缺少严格所有权校验、自动化测试为空。"
        ]
    )

    add_heading(doc, "五、代码结构总图", 1)
    add_table(
        doc,
        ["目录/类群", "职责", "答辩时怎么讲"],
        [
            ["controller", "处理 HTTP 请求、参数获取、页面跳转、会话判断", "相当于系统入口，负责把浏览器请求转成业务调用"],
            ["service", "封装业务规则和模块编排", "用于把控制器和数据库操作解耦，避免 Servlet 直接写 SQL"],
            ["dao", "使用 JDBC 执行 SQL，完成增删改查", "相当于数据库访问层，方便复用和维护"],
            ["entity", "映射用户、摊位、商品、订单、租赁等领域对象", "让业务数据以面向对象方式组织"],
            ["model", "当前主要是 ShoppingCart，会话级购物车模型", "解释为什么购物车不直接频繁落库"],
            ["util", "当前核心是 DatabaseUtil，负责 HikariCP 连接池初始化", "体现基础设施支持"],
            ["sql", "数据库建表与初始化数据脚本", "可以回答数据库设计与初始数据来源"],
            ["webapp/jsp", "前台、摊主后台、管理员后台 JSP 页面", "展示不同角色界面的组织方式"],
        ],
        widths=[Cm(3.2), Cm(5.3), Cm(7)]
    )
    add_heading(doc, "1. 建议重点记住的代码文件", 2)
    add_bullets(
        doc,
        [
            "LoginServlet / RegisterServlet：登录注册与角色跳转入口。",
            "ShoppingCart / CartServlet：购物车状态与优惠券处理。",
            "CheckoutServlet：全项目最关键的下单与拆单逻辑。",
            "OrderServiceImpl / OrderDaoImpl：订单保存、订单明细落库、支付状态更新、取餐号生成。",
            "OwnerOrderServlet：摊主侧订单状态流转。",
            "OwnerProductServlet：商品管理和图片上传。",
            "OwnerLeaseServlet / AdminLeaseServlet / LeaseServiceImpl：租赁申请、审核与合同文本生成。",
            "AdminDashboardServlet：管理员统计面板逻辑。",
            "DatabaseUtil：连接池配置与数据库连接获取。",
        ]
    )

    add_heading(doc, "六、技术选型与原因", 1)
    add_table(
        doc,
        ["技术", "实际用途", "选择原因", "答辩口径"],
        [
            ["Java 17", "后端语言", "类型安全、生态成熟、适合课程项目和分层开发", "选择成熟稳定的 Java 作为后端基础，便于构建面向对象的业务模型"],
            ["Jakarta Servlet 6.0", "控制器层", "直接处理请求响应，链路清晰", "用 Servlet 能完整展示 Java Web 请求处理流程，适合教学与毕设展示"],
            ["JSP + JSTL + EL", "视图层", "部署简单，和 Servlet 配合紧密", "本项目采用传统 MVC，更容易体现前后端协同而不引入额外复杂度"],
            ["JDBC + DAO", "数据访问", "可直接控制 SQL，层次清晰", "有利于展示 SQL、实体映射和数据库设计"],
            ["HikariCP", "连接池", "轻量且性能好", "避免每次请求都创建物理连接，提高数据库访问效率"],
            ["MySQL 8.0", "数据持久化", "关系建模清晰，适合订单类系统", "订单、用户、摊位、租赁等结构化数据适合关系数据库存储"],
            ["Bootstrap 5", "页面样式", "组件成熟、开发效率高", "能在较短开发周期内提供统一的界面样式"],
            ["HttpSession", "登录态与购物车", "会话级状态管理简单直接", "适合单机部署下的小型 Web 系统，便于快速实现登录和购物车"],
        ],
        widths=[Cm(2.2), Cm(3.1), Cm(4.6), Cm(6.1)]
    )

    add_heading(doc, "七、系统架构该怎么讲", 1)
    add_paragraph(
        doc,
        "这个系统本质上是一个典型的 B/S 分层架构项目，浏览器作为客户端，通过 HTTP 请求访问部署在 Servlet 容器中的 Web 应用，应用再通过 DAO 和 JDBC 访问 MySQL 数据库。"
    )
    add_numbered(
        doc,
        [
            "浏览器发起请求，例如用户访问 /login、/cart、/order/checkout 或 /admin/dashboard。",
            "Servlet 读取请求参数、检查会话状态和角色类型，然后调用 Service 层。",
            "Service 层负责组织业务，例如验证账户状态、生成订单、分组拆单、生成合同文本。",
            "DAO 层负责执行 SQL，把 Java 对象映射到 users、orders、leases 等数据表中。",
            "执行结果再返回到 Servlet，由 Servlet 选择转发 JSP 页面或重定向到新的 URL。"
        ]
    )
    add_paragraph(
        doc,
        "如果老师问“为什么要分层”，标准回答是：分层能降低耦合，让控制器只关注请求处理，让 Service 负责业务规则，让 DAO 专注数据库访问，后续维护、测试和替换实现都会更方便。"
    )

    add_heading(doc, "八、数据库设计总览", 1)
    add_paragraph(
        doc,
        "当前业务真正高频使用的是 users、categories、stalls、products、orders、order_items、leases、order_number_sequences、pickup_number_sequences 这些表。除此之外，reviews、complaints、contracts、operation_logs、roles 更像是扩展预留设计，在当前代码中还没有完整打通。"
    )
    add_table(
        doc,
        ["核心表", "关键字段", "作用", "与谁关联"],
        [
            ["users", "username, password_hash, role_type, status", "存放食客/摊主/管理员账户", "被 stalls.owner_id、orders.user_id、leases.owner_id 引用"],
            ["categories", "category_name, region_capacity", "存放摊位品类及区域容量", "被 stalls.category_id 引用"],
            ["stalls", "stall_name, location, status, owner_id, category_id", "存放摊位主体信息", "关联用户、品类、商品、订单、租赁"],
            ["products", "product_name, stall_id, price, status", "存放商品信息", "从属于 stalls，结算时被购物车和订单项引用"],
            ["orders", "order_number, pickup_number, user_id, stall_id, total_amount, status, payment_status", "存放拆单后的主订单", "关联用户、摊位、订单明细"],
            ["order_items", "order_id, product_id, quantity, unit_price, subtotal", "存放订单明细", "从属于 orders，也关联 products"],
            ["leases", "stall_id, owner_id, type, status, start_date, end_date", "存放新租/续租申请和合同文本", "关联用户与摊位"],
            ["order_number_sequences", "date_key, current_seq", "预留按天生成递增订单号", "与订单编号生成逻辑相关"],
            ["pickup_number_sequences", "date_key, current_seq", "按天生成取餐号", "在支付状态更新时使用"],
        ],
        widths=[Cm(3.1), Cm(5.6), Cm(4.2), Cm(4.2)]
    )
    add_heading(doc, "1. 数据库关系怎么说", 2)
    add_bullets(
        doc,
        [
            "一个用户可以是食客、摊主或管理员，由 users.role_type 区分。",
            "一个摊位可以属于一个品类，也可以关联一个摊主。",
            "一个摊位下可以有多个商品。",
            "一个食客一次统一结算后，系统会按摊位拆出多条订单，因此 orders.stall_id 非常关键。",
            "一张订单可以对应多条 order_items，用来表达订单中具体买了哪些商品。",
            "租赁申请通过 leases 表管理，合同文本当前直接以文本形式保存在 contract_content 字段中。"
        ]
    )
    add_heading(doc, "2. 为什么扩展表很多但代码没全部用上", 2)
    add_paragraph(
        doc,
        "这是一个很容易被老师追问的点。当前项目的数据库不仅服务于已经完成的功能，也预留了评价、投诉、合同、操作日志和更细粒度权限控制等扩展能力。也就是说，数据库设计考虑了系统后续演进，但当前代码只完整打通了点餐、订单、租赁和后台治理主线。这个回答既诚实，也能体现你有扩展意识。"
    )

    add_heading(doc, "九、核心业务主线精讲", 1)
    add_heading(doc, "1. 登录、注册与角色控制", 2)
    add_paragraph(
        doc,
        "登录入口在 LoginServlet。控制器读取用户名和密码后调用 UserServiceImpl.login，再由 UserDaoImpl 从 users 表查询账户。当前密码校验采用明文字符串比较，若账户存在且状态为 ACTIVE，则写入 session 的 user 属性并按角色跳转：管理员去 /admin/dashboard，摊主去 /owner/dashboard，食客去 /home。"
    )
    add_paragraph(
        doc,
        "注册入口在 RegisterServlet。它先校验两次密码是否一致，再构造 User 对象。这里有一个非常重要的业务点：如果注册角色是 OWNER，则用户状态默认设为 PENDING，需要管理员后续审核；如果是 DINER，则默认直接 ACTIVE。这个设计体现了平台对摊主入驻的审核思路。"
    )
    add_bullets(
        doc,
        [
            "优点：实现简单，能清晰体现注册、审核、登录和角色跳转的业务逻辑。",
            "不足：密码未加密；权限校验分散在各个 Servlet 中，没有集中式过滤器或拦截器。",
            "答辩建议：先承认这是课程级实现，再补充生产环境会引入 BCrypt 和统一鉴权组件。"
        ]
    )

    add_heading(doc, "2. 首页和摊位浏览", 2)
    add_paragraph(
        doc,
        "HomeServlet 会调用 ProductServiceImpl.getLatestAvailableProducts(6) 和 StallServiceImpl.getLatestOpenStalls(4)，把最新可售商品和营业中的推荐摊位展示在首页。StallDetailServlet 则根据摊位 id 查询摊位详情，并且只把状态为 AVAILABLE 的商品展示给食客。这个设计能体现“前台只暴露可营业、可售卖资源”的业务约束。"
    )

    add_heading(doc, "3. 购物车为什么放在 Session 中", 2)
    add_paragraph(
        doc,
        "ShoppingCart 是一个会话级模型，内部使用 Map<Integer, CartItem> 以商品 id 作为键存储购物车项，提供 addItem、updateQuantity、removeItem、clear、getTotalAmount 等方法。CartServlet 会在用户 Session 中读取或创建 cart 对象，因此加购、改数量、删商品和清空购物车都发生在会话内存中，而不是每次都写数据库。"
    )
    add_bullets(
        doc,
        [
            "这样设计的原因是购物车变动非常频繁，频繁落库会增加不必要的数据库读写开销。",
            "Session 方案适合单机部署、开发成本低，特别适合课程项目或毕设原型。",
            "局限也很明显：多机部署下 Session 共享困难，服务器重启后购物车会丢失，不适合高并发生产环境。"
        ]
    )

    add_heading(doc, "4. 优惠券逻辑怎么实现", 2)
    add_paragraph(
        doc,
        "CartServlet 内部维护了一个 couponMap，当前硬编码了 FC10、NIGHT12、WELCOME5 三种优惠券及其减免金额。用户提交 couponCode 后，系统会校验券码是否存在、购物车是否为空、优惠金额是否超过当前总价，并把最终折扣额和券码存入 session 的 cartDiscount 与 cartCoupon。"
    )
    add_paragraph(
        doc,
        "如果用户后续修改购物车内容，adjustDiscount 方法会重新检查折扣是否仍然合理，例如当购物车总价低于已有折扣时，会自动把折扣压缩到新的购物车总价。这个细节说明当前实现虽然简单，但考虑了基本的一致性。"
    )

    add_heading(doc, "5. 最核心的功能：统一结算与自动拆单", 2)
    add_paragraph(
        doc,
        "CheckoutServlet 是全项目最值得重点讲解的类。它先从 session 中取出当前用户和购物车，如果购物车为空则不能结算。然后通过 Java Stream 的 Collectors.groupingBy，按照每个购物车商品所属的 stallId 对购物车项分组。这样一个购物车就会被拆成“摊位 A 的商品集合、摊位 B 的商品集合、摊位 C 的商品集合”等多个分组。"
    )
    add_numbered(
        doc,
        [
            "读取用户、购物车、支付方式、优惠券信息。",
            "按 stallId 对购物车项进行分组，形成多个摊位订单分组。",
            "计算购物车总额和可用优惠金额。",
            "如果使用优惠券，则按各摊位订单金额占总金额的比例分摊折扣，最后一个分组吃掉剩余折扣，避免四舍五入残差。",
            "针对每个摊位创建一条 Order 主记录，并生成对应的 OrderItem 列表。",
            "调用 OrderServiceImpl.createOrder 保存订单和明细。",
            "随后调用 updatePaymentStatus，把支付状态更新为 PAID，并生成取餐号。",
            "全部完成后清空购物车、清除优惠券会话数据，并跳转到食客订单列表。"
        ]
    )
    add_paragraph(
        doc,
        "这条链路的业务意义在于：食客视角是一次统一结算，但摊主视角必须是分摊位独立接单，因为每个摊主只负责自己的商品。自动拆单正是把“消费端的一体化体验”和“经营端的摊位独立性”连接起来的关键设计。"
    )

    add_heading(doc, "6. 订单号、取餐号与序列表的设计", 2)
    add_paragraph(
        doc,
        "OrderServiceImpl.createOrder 在插入订单前会先生成一个 TMP + UUID 截断的临时订单号，用于保证订单插入时 order_number 字段非空且唯一。随后 OrderDaoImpl.updatePaymentStatus 在支付状态变为 PAID 时，会对当前订单行执行 SELECT ... FOR UPDATE，并尝试为订单补充分配正式编号和取餐号。取餐号来自 pickup_number_sequences 序列表，格式类似 P202601300001。"
    )
    add_paragraph(
        doc,
        "这里要特别注意一个答辩细节：按照当前代码路径，订单主表在创建时已经带有 TMP 临时订单号，后续支付状态更新时会保留这个临时号，因此 order_number_sequences 的正式订单号生成逻辑实际上没有完全接通，真正稳定使用的是取餐号序列表。这个点不要回避，直接说明“设计上考虑了正式订单号和取餐号双序列，但当前版本只把取餐号完整打通，订单号后续需要在支付成功时统一重生成为正式流水号”。老师通常会认可这种诚实且具体的回答。"
    )
    add_bullets(
        doc,
        [
            "优点：使用 FOR UPDATE 和单日序列表，可以避免并发下重复取餐号。",
            "不足：正式订单号生成链路未彻底闭环，当前保留了 TMP 临时号。",
            "改进：把 order_number 的正式生成迁移到支付成功后统一执行，或者在创建订单时允许先为空。"
        ]
    )

    add_heading(doc, "7. 食客端订单查询", 2)
    add_paragraph(
        doc,
        "DinerOrderServlet 根据 session 中当前食客 id 查询自己的订单列表。OrderDetailServlet 则根据订单 id 查询主订单和订单明细，并在食客角色下校验该订单是否属于自己。这里虽然没有使用复杂权限框架，但至少体现了“查看他人订单需要被禁止”的基本安全意识。"
    )

    add_heading(doc, "8. 摊主侧：经营数据、商品管理、订单流转", 2)
    add_paragraph(
        doc,
        "OwnerDashboardServlet 会统计摊主名下所有摊位的待处理订单数、今日营收、昨日营收对比和商品总量。统计方式是逐个摊位拉取订单，再根据状态和下单日期做汇总。虽然是偏朴素的实现，但能体现数据面板如何基于业务数据实时计算。"
    )
    add_paragraph(
        doc,
        "OwnerProductServlet 负责商品增删改查，并在进入或提交时校验当前摊位是否属于当前摊主。图片上传通过 MultipartConfig 和 Part 实现，图片会复制到运行时 uploads 目录以及项目 uploads 目录中，再以 /uploads/文件名 的方式对外访问。这个模块适合用来回答“文件上传怎么做”和“如何防止摊主管理别人的商品”。"
    )
    add_paragraph(
        doc,
        "OwnerOrderServlet 则负责摊主查看订单并推进状态流转，支持 confirm、prepare、complete、cancel 四种动作，对应 PENDING -> CONFIRMED -> PREPARING -> COMPLETED 或 CANCELLED。这里有一个需要主动说明的风险点：GET 逻辑是通过摊主名下摊位拉订单，POST 更新时虽然做了角色校验，但代码注释里明确写了没有做严格的订单所有权校验，因此这是当前版本的安全短板之一。"
    )

    add_heading(doc, "9. 摊位租赁申请与合同生成", 2)
    add_paragraph(
        doc,
        "OwnerLeaseServlet 支持摊主查看自己的租赁记录，并提交新租或续租申请。提交时会写入 leases 表，状态默认是 PENDING。管理员在 AdminLeaseServlet 中可以查看所有租赁记录、按关键字搜索、更新审核状态，以及调用 LeaseServiceImpl.generateContract 生成合同文本并回写到 contract_content 字段。"
    )
    add_paragraph(
        doc,
        "这部分业务体现了项目并不只是一个点餐系统，而是把美食街管理中的摊位租赁和平台治理也纳入了进来。答辩时如果老师问“管理端除了看订单还能做什么”，你就可以重点讲用户审核、品类容量、租赁审批和合同生成。"
    )

    add_heading(doc, "10. 管理员后台：治理逻辑怎么讲", 2)
    add_paragraph(
        doc,
        "AdminDashboardServlet 会计算摊位出租率、订单支付率和热门摊位 Top5。其核心做法是先拉全量摊位和订单，再用 Java Stream 在内存中统计。AdminUserServlet 负责用户状态管理，AdminStallServlet 负责摊位列表查询，AdminOrderServlet 负责订单监控，AdminCategoryServlet 可以维护 categories.region_capacity 字段，AdminLeaseServlet 则处理租赁审核和合同生成。"
    )
    add_bullets(
        doc,
        [
            "摊位出租率 = owner_id 非空的摊位数 / 摊位总数。",
            "订单支付率 = PAID 订单数 / 全部订单数。",
            "热门摊位 = 按订单数量分组排序后的前 5 名。",
            "品类容量 region_capacity 可以理解为某个区域或品类可承载的摊位上限，属于平台治理参数。"
        ]
    )

    add_heading(doc, "十、项目亮点总结", 1)
    add_bullets(
        doc,
        [
            "业务闭环完整：不是单纯商品展示，而是覆盖食客消费、摊主经营、管理员治理三条线。",
            "核心业务有辨识度：跨摊位购物车统一结算并自动拆单，这一点比普通单商户点餐系统更有特点。",
            "分层结构清晰：Servlet、Service、DAO、Entity 的职责边界比较明确，便于答辩讲清楚。",
            "数据库设计具备扩展意识：除了已完成功能，还预留了评价、投诉、合同、日志、角色权限等扩展表。",
            "取餐号并发控制思路正确：通过每日序列表和 FOR UPDATE 实现递增编号生成。",
            "管理功能不只是 CRUD：加入了出租率、支付率、热门摊位、租赁审核、合同文本等治理维度。"
        ]
    )

    add_heading(doc, "十一、当前不足与老师最爱追问的风险点", 1)
    add_table(
        doc,
        ["问题点", "当前表现", "为什么会被老师问", "标准改进口径"],
        [
            ["密码安全", "登录与注册使用明文密码比较和存储", "任何账户系统都绕不开安全问题", "课程版为了聚焦业务流程而简化，生产环境必须改为 BCrypt/Scrypt 哈希并配合盐值"],
            ["统一鉴权", "每个 Servlet 手动判断 session 和 roleType", "老师会问为什么不用 Filter/Interceptor", "当前实现便于展示逻辑，但下一步应抽取统一认证过滤器和权限中间件"],
            ["订单事务", "save(order) 与 saveOrderItems(items) 不在同一数据库事务中", "老师会问主表成功明细失败怎么办", "应使用同一连接显式事务或引入 Spring 事务管理，保证原子性"],
            ["摊主越权", "OwnerOrderServlet POST 未做严格订单所有权校验", "老师会问是否能伪造订单 id 修改他人订单", "应在更新前根据 orderId 反查 stallId 和 ownerId 完成服务端校验"],
            ["订单号闭环", "当前主要使用 TMP 临时订单号，正式序列号未完全打通", "老师会问为什么数据库里订单号不是流水号", "应在支付成功时统一生成正式订单号，或者把订单创建与支付确认拆成两个阶段"],
            ["购物车存储", "Session 内存保存，服务重启或多机部署下不稳定", "老师会问扩展到高并发怎么办", "可迁移到 Redis 或数据库，并增加购物车持久化和过期策略"],
            ["测试覆盖", "src/test/java 为空，缺少自动化测试", "老师会问如何验证系统质量", "当前主要做了功能级手工测试，后续可补 Service/DAO 单元测试与集成测试"],
            ["扩展表落地", "reviews/complaints/contracts/roles 等表未完全接入代码", "老师会问是否是堆砌设计", "这些表是为后续评价、投诉、细粒度权限和正式合同管理预留的扩展能力"],
        ],
        widths=[Cm(2.7), Cm(4.5), Cm(4.3), Cm(5)]
    )
    add_paragraph(
        doc,
        "答辩时遇到不足类问题，切记不要硬拗。正确姿势是：先承认这是当前版本的真实边界，再解释为什么当时这样取舍，最后给出可执行的改进方案。这样比一味强调“我这个系统没有问题”更专业。"
    )

    add_heading(doc, "十二、演示顺序与复习清单", 1)
    add_heading(doc, "1. 最稳的现场演示顺序", 2)
    add_numbered(
        doc,
        [
            "先用一句话介绍项目定位和三类角色。",
            "登录食客账号，进入首页，展示推荐摊位和商品。",
            "进入两个不同摊位的详情页，各加入一个商品到购物车。",
            "打开购物车，修改数量，输入优惠券，展示总价变化。",
            "执行结算，强调系统会按 stallId 自动拆单。",
            "进入食客订单列表，展示虽然一次结算，但出现了多条摊位订单。",
            "登录摊主账号，进入订单管理页，演示接单、备餐、完成等状态流转。",
            "进入摊主商品管理，展示本摊位的商品增删改和图片上传。",
            "登录管理员账号，展示出租率、支付率、热门摊位、用户管理、品类容量、租赁审核。",
            "最后总结亮点和不足，主动说出后续优化方向。"
        ]
    )
    add_heading(doc, "2. 复习时一定要能脱口而出的点", 2)
    add_bullets(
        doc,
        [
            "为什么需要三类角色，而不是只做用户和管理员。",
            "为什么购物车要放在 Session 中。",
            "CheckoutServlet 为什么是核心类，它到底做了哪几步。",
            "自动拆单的依据是什么，为什么按 stallId 拆。",
            "优惠券为什么要按金额比例分摊，而不是平均分。",
            "HikariCP 的作用是什么。",
            "管理员统计页的三项指标怎么计算。",
            "当前版本的明显不足有哪些，你准备怎么改。"
        ]
    )
    add_heading(doc, "3. 一天内冲刺复习建议", 2)
    add_numbered(
        doc,
        [
            "先看本指南的“项目速览”和“核心业务主线”，建立整体框架。",
            "再打开 CheckoutServlet、ShoppingCart、OrderDaoImpl 三个类，彻底搞懂点餐主线。",
            "再复习 LoginServlet、RegisterServlet、OwnerOrderServlet、AdminDashboardServlet、LeaseServiceImpl。",
            "最后背 30 道高频追问，重点背标准口径而不是逐字死记。"
        ]
    )

    add_heading(doc, "十三、答辩万能回答模板", 1)
    add_paragraph(
        doc,
        "遇到老师追问时，建议统一用“四句结构”回答，这样条理最清楚，也不容易慌。"
    )
    add_numbered(
        doc,
        [
            "先给结论：直接回答“是怎么做的”或“为什么这么设计”。",
            "再讲原理：说明背后的业务逻辑或工程逻辑。",
            "再落到本项目：指出具体类、字段、表或流程。",
            "最后补边界：承认当前不足，并给出下一步优化方案。"
        ]
    )
    add_paragraph(
        doc,
        "例如老师问“为什么购物车放在 Session 里”，你不要只说“方便”，而要说：因为购物车操作频繁且是强会话场景，当前单机教学系统为了减少数据库压力采用 Session 保存；具体由 ShoppingCart 和 CartServlet 实现；但在分布式环境下会迁移到 Redis 做共享和持久化。"
    )

    doc.add_page_break()
    add_heading(doc, "十四、老师高频追问 30 题标准答题卡", 1)
    add_paragraph(
        doc,
        "下面 30 题按答辩真实场景整理，回答口径尽量做到“有结论、有代码依据、有边界说明”。如果老师问法不同，也可以把其中的句型灵活组合。"
    )

    questions = [
        (
            "请你用一句话概括这个项目。",
            "这是一个基于 Java Web 的多角色美食街摊位管理系统，面向食客、摊主和管理员三类用户，覆盖浏览商品、购物车、统一结算、按摊位自动拆单、摊主接单以及管理员治理的完整流程。",
            ["如果需要展开，可以再补一句：项目最大的业务特点是一次结算、多摊位拆单。"]
        ),
        (
            "这个项目主要解决了什么痛点？",
            "它主要解决了传统美食街场景下跨摊位消费体验差、摊主人工接单效率低、管理员缺少整体经营数据三类问题。食客可以统一下单，摊主可以独立处理分配给自己的订单，管理员可以统一查看摊位、用户、订单和租赁信息。",
            None
        ),
        (
            "为什么要设计食客、摊主、管理员三类角色？",
            "因为这三类用户在业务中的目标完全不同。食客关注浏览和下单，摊主关注商品和订单处理，管理员关注平台治理和审核。如果只做普通用户和管理员，会把经营侧的摊主职责混在一起，业务边界不清晰，也无法体现多主体协同的特点。",
            None
        ),
        (
            "为什么技术上选择 Servlet + JSP，而不是 Spring Boot + 前后端分离？",
            "本项目是课程/毕设场景，目标是把 Java Web 的请求处理链路、分层思想和数据库交互过程讲清楚。Servlet + JSP 的技术栈更直接，能完整展示从请求进入到 JSP 渲染返回的全过程，复杂度更低，也更适合在答辩中解释具体实现。当然如果面向更完整的工程化版本，后续完全可以迁移到 Spring Boot 加 Vue 或 React 的架构。",
            None
        ),
        (
            "系统的分层架构是怎样的？",
            "系统采用 controller、service、dao、entity 的典型分层。controller 负责处理请求和页面跳转，service 负责组织业务规则，dao 负责 JDBC 和 SQL 执行，entity 负责领域对象表达。这样的分层有利于降低耦合，让页面逻辑、业务逻辑和数据库逻辑分离。",
            None
        ),
        (
            "为什么要把 DAO 和 Service 分开？Servlet 里直接写 SQL 不行吗？",
            "直接在 Servlet 里写 SQL 虽然能实现功能，但控制器会同时承担请求处理、业务判断和数据库访问三种职责，后续维护会很混乱。DAO 和 Service 分开后，Servlet 更像路由入口，Service 负责业务编排，DAO 负责数据访问，这样结构更清晰，也更利于后续测试和扩展。",
            None
        ),
        (
            "为什么数据库选 MySQL？",
            "因为这个项目的核心数据都是强结构化关系数据，比如用户、摊位、商品、订单、订单明细和租赁记录，天然适合关系数据库建模。MySQL 成熟稳定、开发部署门槛低，也支持外键、索引、事务和枚举字段，适合订单类系统的课程实现。",
            None
        ),
        (
            "为什么要用 HikariCP 连接池？",
            "如果每次请求都临时创建和销毁 JDBC 连接，数据库开销会比较大。HikariCP 可以复用数据库连接，减少创建成本，提高响应效率。当前 DatabaseUtil 在类加载时读取 db.properties，初始化最大连接数、最小空闲连接数和超时时间，为 DAO 层统一提供连接获取入口。",
            None
        ),
        (
            "为什么购物车放在 Session 中，而不是数据库里？",
            "因为购物车是一个强会话、强交互、频繁变动的数据结构，加购、减购、改数量都很高频。如果每次操作都落数据库，代价会比较高。当前项目是单机课程版实现，所以用 Session 保存购物车最直接，ShoppingCart 负责状态结构，CartServlet 负责会话读写。不过在分布式部署下，这种方式需要升级为 Redis 或数据库持久化方案。",
            None
        ),
        (
            "优惠券是怎么实现的？",
            "当前优惠券逻辑实现得比较轻量，CartServlet 内部维护了一个 couponMap，预置了 FC10、NIGHT12、WELCOME5 三种券码。系统会校验券码是否有效、购物车是否为空、折扣是否超过总价，并把最终折扣额和券码放进 session。这样可以快速打通优惠逻辑，但生产环境中优惠券更适合存数据库并配合使用次数、有效期和用户范围做控制。",
            None
        ),
        (
            "自动拆单是怎么做的？",
            "核心逻辑在 CheckoutServlet。系统先读取购物车中的所有商品，再根据每个商品所属的 stallId 使用 groupingBy 分组。每个分组代表一个摊位的商品集合，随后系统会为每个摊位分别创建一条 Order 主记录和若干 OrderItem 明细。这样食客只做一次结算，但不同摊主可以看到属于自己的独立订单。",
            None
        ),
        (
            "为什么按摊位拆单，而不是按商品或者按用户拆单？",
            "因为摊位是经营主体，订单最终要由摊主接单、备餐和结算，所以拆单粒度应该围绕摊位，而不是商品。按商品拆单会把同一摊位的一组商品拆得太碎，按用户拆单又无法区分多个摊主的处理边界，因此按 stallId 拆是最符合业务流程的。",
            None
        ),
        (
            "如果用了优惠券，为什么要按金额比例分摊到每个子订单？",
            "因为食客是对整张购物车使用优惠券，但系统最终要拆成多张摊位订单。为了保证每个摊位订单的金额与原始商品金额比例一致，就要按各分组金额占购物车总额的比例分摊折扣。这样比平均分更公平，也更符合财务核算逻辑。",
            None
        ),
        (
            "为什么系统里会出现 TMP 开头的订单号？",
            "这是当前版本的设计取舍。OrderServiceImpl.createOrder 在订单插入前先生成一个 TMP 加 UUID 截断的临时订单号，目的是保证 order_number 字段先满足非空且唯一。后续支付状态更新时，系统会生成正式取餐号，但正式订单号的序列化逻辑目前没有完全闭环，所以数据库里仍然可能保留 TMP 号。这个点我会作为后续优化项，在支付成功时统一生成正式订单号。",
            None
        ),
        (
            "系统为什么还设计了 order_number_sequences 和 pickup_number_sequences 两张序列表？",
            "因为系统希望支持按天递增的业务编号。取餐号已经通过 pickup_number_sequences 实际使用，用于生成类似 P202601300001 的号码；order_number_sequences 则是为正式订单流水号预留的序列机制，只是当前版本还没有完全打通。这说明项目在设计上考虑了并发编号问题，但代码实现还有收尾空间。",
            None
        ),
        (
            "如何保证取餐号在并发场景下不重复？",
            "核心做法是在 OrderDaoImpl.updatePaymentStatus 中使用 SELECT ... FOR UPDATE 锁住当前订单和当天的序列表记录，再在同一事务里读取 current_seq、加一并更新。这样在并发下，同一时刻只有一个事务可以修改当天的序列值，所以取餐号可以保持唯一递增。",
            None
        ),
        (
            "系统的权限控制是怎么做的？",
            "当前权限控制主要是基于 Session 中的 user 对象和 users.role_type 字段完成的。各个 Servlet 在 doGet 或 doPost 里会判断是否登录、角色是否匹配，例如管理员页面只允许 ADMIN 访问，摊主页面只允许 OWNER 访问。它的优点是直观，缺点是校验逻辑比较分散，下一步更好的做法是抽取统一的 Filter。 ",
            None
        ),
        (
            "为什么摊主注册后默认是 PENDING，而食客是 ACTIVE？",
            "因为摊主会涉及摊位经营、商品发布和租赁申请，平台需要先审核其资格，所以 OWNER 注册后会被设置为 PENDING。食客只涉及消费行为，风险相对低，所以可以默认 ACTIVE。这体现了不同业务角色在准入机制上的差异。 ",
            None
        ),
        (
            "系统如何防止摊主管理别人的商品？",
            "OwnerProductServlet 会先读取 stallId，再通过 StallService 查询该摊位，并校验该摊位 ownerId 是否等于当前登录摊主 id。如果不匹配，系统直接返回 403。也就是说，商品管理权限不是只看角色，而是同时看角色和资源归属。 ",
            None
        ),
        (
            "图片上传功能是怎么实现的？",
            "OwnerProductServlet 使用 @MultipartConfig 支持 multipart/form-data 请求，再通过 Part 读取上传文件。系统会检查文件是否为图片类型，生成随机文件名，把图片复制到 uploads 目录，并把访问路径保存到 product.imageUrl。这个实现满足课程项目需求，但在生产环境中还需要增加文件大小、后缀、木马脚本和对象存储方面的安全控制。 ",
            None
        ),
        (
            "管理员统计大盘里的数据是怎么来的？",
            "AdminDashboardServlet 会先取出全部摊位和全部订单，然后在内存中计算三个指标：摊位出租率、订单支付率和热门摊位 Top5。出租率的分子是 owner_id 不为空的摊位数，支付率的分子是 payment_status 为 PAID 的订单数，热门摊位则按 stallName 分组统计订单数量后排序得到。 ",
            None
        ),
        (
            "租赁申请和合同生成这条业务链是怎样的？",
            "摊主在 OwnerLeaseServlet 中提交新租或续租申请，记录会写入 leases 表并默认设置为 PENDING。管理员在 AdminLeaseServlet 中查看申请后，可以修改状态为 APPROVED 或 REJECTED；如果需要生成合同，则调用 LeaseServiceImpl.generateContract，根据摊位、摊主、起止日期拼接合同文本，再回写到 contract_content 字段。 ",
            None
        ),
        (
            "categories 表里的 region_capacity 字段是什么意思？为什么需要它？",
            "它可以理解为某类区域或某类品类可容纳的摊位容量上限，属于平台治理参数，而不是食客侧消费数据。管理员可以在 AdminCategoryServlet 中维护这个值，用来支持后续的摊位规划和品类布局管理。这个字段体现了系统不仅关注点餐，也关注经营资源分配。 ",
            None
        ),
        (
            "为什么数据库里设计了 reviews、complaints、contracts、roles 等表，但代码没有全部用上？",
            "这些表属于扩展预留设计。当前代码已经重点打通了订单、摊位、商品、租赁和后台治理主线，而评价、投诉、正式合同管理和细粒度权限控制是后续准备继续扩展的能力。也就是说，数据库设计先考虑了系统完整演进方向，但开发优先级放在了核心业务闭环上。 ",
            None
        ),
        (
            "你认为这个项目当前最大的安全问题是什么？",
            "最大的问题有两个：一是密码目前仍然明文存储和比较，二是权限校验分散且部分接口的资源归属校验不够严格，例如摊主更新订单状态时没有做完整的订单归属验证。前者需要改为 BCrypt 哈希，后者需要补统一 Filter 和服务端二次校验。 ",
            None
        ),
        (
            "如果订单主表插入成功，但订单明细插入失败，会发生什么？",
            "当前版本会出现数据不一致，因为 OrderServiceImpl.createOrder 先调用 orderDao.save，再调用 orderDao.saveOrderItems，这两步没有绑定在同一数据库事务里。如果中间失败，就可能出现孤立订单主记录。这是一个典型的事务性问题，后续应该通过同一连接的手工事务或 Spring 声明式事务来保证原子性。 ",
            None
        ),
        (
            "这个项目在并发和扩展性方面还有哪些限制？",
            "一方面，购物车放在 Session 中，天然更适合单机环境；另一方面，很多统计和搜索是先拉全量数据再在内存中过滤，数据量一大性能会下降。另外，部分权限校验和事务控制还停留在课程版实现，不适合直接部署到高并发生产环境。后续可通过 Redis、分页查询、SQL 聚合、统一鉴权和事务重构来提升。 ",
            None
        ),
        (
            "这个项目是怎么测试的？测试还有什么不足？",
            "当前项目更偏功能验证，主要通过手工场景测试来检查注册登录、购物车、下单拆单、订单流转、租赁审核和后台统计是否正确。源码中的 src/test/java 目前为空，说明自动化测试覆盖还没有建立起来。后续可以补充 Service 层单元测试、DAO 层集成测试以及关键业务链路的接口测试。 ",
            None
        ),
        (
            "你觉得这个项目最有价值的创新点或者亮点是什么？",
            "我认为最有价值的点不是单独的 CRUD，而是把多角色和多摊位场景串成了一个完整闭环，尤其是跨摊位统一结算和按 stallId 自动拆单这条链路。它既照顾了食客的统一支付体验，也满足了摊主独立接单的业务边界，是这个系统区别于普通单店点餐系统的关键。 ",
            None
        ),
        (
            "如果让你把这个项目往生产级方向升级，最先做哪三件事？",
            "第一，补安全：密码哈希、统一认证过滤器、接口级资源归属校验。第二，补一致性：把订单创建、明细插入、支付状态更新放进完整事务中。第三，补可扩展性：把购物车迁移到 Redis，引入分页与聚合 SQL，并逐步迁移到更成熟的 Spring Boot 架构。这样能先把最关键的安全性、一致性和扩展性补齐。 ",
            None
        ),
        (
            "你接下来准备怎么继续完善这个项目？",
            "后续我会从三个方向继续做：第一是把现有课程版实现补成更完整的工程化版本，重点解决安全、事务和鉴权问题；第二是把评价、投诉、正式合同管理和操作日志等预留模块真正接入业务；第三是优化前端交互和统计查询方式，让系统在更多真实使用场景下可用。这样项目就能从“能演示”逐步走向“能稳定运行”。 ",
            None
        ),
    ]

    for idx, (question, answer, points) in enumerate(questions, start=1):
        add_question_card(doc, idx, question, answer, points)

    add_heading(doc, "十五、最后的答辩提醒", 1)
    add_bullets(
        doc,
        [
            "不要把项目说成自己没有实现的技术栈。当前真实技术栈是 Servlet + JSP，不是 Spring Boot + Vue。",
            "不要回避短板。越是容易被老师抓住的点，越要提前准备“当前做法 + 风险 + 改进”。",
            "自动拆单一定要重点讲，因为这是项目最有识别度的业务能力。",
            "如果老师突然让你讲数据库，不要从表名堆砌开始，要从 users、stalls、products、orders、order_items、leases 六张核心表开始讲。",
            "如果现场紧张，就用“四句结构”回答：结论、原理、本项目实现、后续优化。"
        ]
    )

    doc.save(OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
